import pandas as pd
import matplotlib.pyplot as plt
from docx import Document


def df_to_word_table(document, df, title):
    # Sets up word document
    document.add_heading(title, level=2)
    table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1] + 1)
    table.style = "Table Grid"

    # Add column headers
    table.cell(0, 0).text = ""
    for j, col_name in enumerate(df.columns, 1):
        table.cell(0, j).text = str(col_name)

    # Add row headers and data
    for i, row_name in enumerate(df.index, 1):
        table.cell(i, 0).text = str(row_name)
        for j, value in enumerate(df.iloc[i - 1], 1):
            table.cell(i, j).text = f"{value:.4f}" if isinstance(value, (float, int)) else str(value)


def main():
    prices = pd.read_csv("modified-prices-split-adjusted.csv")

    # Plot histograms
    plt.rc('font', size=14)
    plt.rc('axes', labelsize=14, titlesize=14)
    plt.rc('legend', fontsize=14)
    plt.rc('xtick', labelsize=10)
    plt.rc('ytick', labelsize=10)
    prices.hist(bins=50, figsize=(12, 8))
    plt.show()

    # Calculate correlation matrix and summary statistics
    corr_matrix = prices.corr(numeric_only=True)
    summary_stats = prices.describe()

    # Export to Word document
    document = Document()
    document.add_heading("Stock Data Analysis", level=1)

    df_to_word_table(document, corr_matrix, "Correlation Matrix")
    df_to_word_table(document, summary_stats, "Summary Statistics")

    output_filename = "stock_analysis_report.docx"
    document.save(output_filename)
    print(f"Report saved as {output_filename}")


if __name__ == "__main__":
    main()
